#############################
### FUNCTIONS DEFINITIONS ###
#############################

# Imports

import pandas as pd
import docx 
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from constants import *


def add_hyperlink(paragraph, url, text):
    
    '''
    Function that allows to embed an hyperlink in a paragraph. 
    
    credits: @johanvandegriff https://github.com/python-openxml/python-docx/issues/74

    params: paragraph (docx.paragraph, paragraph to which add the hyperlink).
            url (string, url of the hyperlink).
            text (string, text to display hyperlinked).
    return: hyperlink (hyperlinked paragraph). 
    '''
    
    # This gets access to the document.xml.rels file and gets a new relation id value.
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values.
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element.
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element.
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element.
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


def create_save_company_profiles(df_raw_data, df_sections, intro_for_analysis_docx, sections_headings, sections_dict, root, sub_folders_names):
    
    '''
    TODO  
    params: .
            .
    return: . 
    '''
    
    for company in df_raw_data.index:
        document = docx.Document()
        document.add_heading('{}: trasparenza e anti-corruzione.'.format (company), 1)
        document.add_heading('Analisi a cura di Transparency International Italia'.format (company), 3)
        document.add_paragraph('')
        document.add_paragraph(intro_for_analysis_docx.format (company, 
                                                               df_sections.loc[company, "Bands"], \
                                                               df_sections.loc[company, "TRAC_Index"])) \
                                                               .alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        document.add_paragraph('---')

        for key, value in sections_dict.items():

            # Add heading of the section, style 4. 
            document.add_heading("Sezione {}: {}".format(key, sections_headings[int(key)-1]), 4) 

            # Handles 0.00 cases (just get the answer of the first question of the section - all answers are the same)
            if  df_sections.loc[company, "Section_{}".format(key)] == 0.00:
                document.add_paragraph("""Alla sezione {}, {} ha ottenuto un punteggio pari a {:.1%} perché {}"""
                                       .format(key, 
                                               company, 
                                               df_sections.loc[company, "Section_{}".format(key)],
                                               df_raw_data.loc[company, "Comment_{}_1".format(key)]))\
                                       .alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
                    
            # Handles None cases (just get the answer of the first question of the section all answers are the same)
            elif df_sections.loc[company, "Section_{}".format(key)] is None:
                document.add_paragraph("""La sezione {}, è stata disattivata (è stato attributio un NA, non applicabile)
                perché {}"""
                .format(key, 
                        df_raw_data.loc[company, "Comment_{}_1".format(key)]))\
                .alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # All other cases (add section summary with the average of that section for the company and print space)

            else:
                document.add_paragraph("""Alla sezione {}, {} ha ottenuto un punteggio pari a {:.1%}"""
                                       .format(key, 
                                               company, 
                                               df_sections.loc[company, "Section_{}".format(key)]))\
                                       .alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Insert paragraph for every question. 
                for question in range(1, len(value)+1):
                    p = document.add_paragraph("{} ha ottenuto un punteggio pari a {} alla domanda {}, perché {} Si veda qui: "
                                               .format (company, 
                                                        df_raw_data.loc[company, "Score_{}_{}".format(key, question)],
                                                        "{}_{}".format(key, question), 
                                                        df_raw_data.loc[company, "Comment_{}_{}".format(key, question)]), 
                                               style = 'List Bullet')
                    add_hyperlink(p, 
                                  df_raw_data.loc[company, "URL_{}_{}".format(key, question)], 
                                  df_raw_data.loc[company, "Source_{}_{}".format(key, question)])
                    
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
        # Break page. 
        document.add_page_break()
        
        # Store file at the right path.
        path = root + "/" + sub_folders_names[1] + "/" + "{}".format(company)
        document.save(path + "/" + "{}_analisi".format(company) + ".docx")
        
    return None 


    

